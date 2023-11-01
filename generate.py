from docx import Document
from os import listdir
from bs4 import BeautifulSoup
from os.path import exists
from os import mkdir
from docx.shared import RGBColor, Pt

# maps a specific html tag to a document text style :3
document_actions = {
  "h1": lambda document, text : document.add_heading(text, 1),
  "h2": lambda document, text : document.add_heading(text, 2),
  "p": lambda document, text : document.add_paragraph(text),
  "a": lambda document, text : document.add_paragraph(text),
  "li": lambda document, text : document.add_paragraph(text, style='List Bullet'),
}

class DocumentMetadata:
  def __init__(self, title=None, url=None, filepath=None, onedrivelink=None):
    self.title = title
    self.url = url
    self.onedrivelink = onedrivelink
    self.filepath = filepath


def generate_documents(dirname, documents_folder, replace=True):
  if not exists(documents_folder):
    mkdir(documents_folder)
    
  metadata = dict()
    
  for subdirname in listdir(dirname):
    metadata[subdirname] = []
    if not exists(f'{documents_folder}/{subdirname}'):
      mkdir(f'{documents_folder}/{subdirname}')
    for filename in listdir(f'{dirname}/{subdirname}'):
      document_metadata = generate_document(
        f'{dirname}/{subdirname}/{filename}',
        documents_folder,
        replace
      )
      metadata[subdirname].append(document_metadata)
  
  return metadata

def generate_document(filepath, dirname, replace) -> DocumentMetadata:
  document = None
  document_metadata = DocumentMetadata()
  doc_filepath = filepath.replace('.html', '.docx').replace('./articles', dirname)
  save_to_file = replace or (not replace and not exists(doc_filepath))
  if save_to_file:
    document = Document()
    styles = document.styles
    styles['Normal'].font.name = 'Rubik'
    styles['Heading 1'].font.name = 'Rubik Light 300'
    styles['Heading 1'].font.size = Pt(20)
    styles['Heading 1'].font.color.rgb = RGBColor.from_string('000000')
    styles['Heading 2'].font.name = 'Rubik Light 300'
    styles['Heading 2'].font.size = Pt(14)
    styles['Heading 2'].font.color.rgb = RGBColor.from_string('000000')

  with open(filepath, 'rb') as file:
    print('generating word file', doc_filepath)
    
    soup = BeautifulSoup(file.read(), 'html.parser')
    article_url = soup.find(id='article_url')
    document_metadata.url = article_url['href']
    heading = soup.find(class_='mw-page-title-main')
    document_metadata.title = heading.getText()
    
    if save_to_file:
      document.add_paragraph(f'Article Link: {document_metadata.url}')
      document.add_heading(document_metadata.title, 1)
    
      content = soup.find(class_='mw-parser-output')
      for child in content.children:
        if child.name == 'ul':
          for grandchild in child.children:
            document_action = document_actions.get(grandchild.name, None)
          if document_action == None:
            continue
          document_action(document, child.getText())
        else:
          document_action = document_actions.get(child.name, None)
          if document_action == None:
            continue
          document_action(document, child.getText())
  
  if save_to_file:
    document.save(doc_filepath)  
  document_metadata.filepath = doc_filepath
  return document_metadata